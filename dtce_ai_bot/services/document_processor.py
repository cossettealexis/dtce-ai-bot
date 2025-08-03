import asyncio
from typing import Optional, Dict, Any
import docx
import PyPDF2
import openpyxl
import io
import structlog

from src.models import DocumentMetadata

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """Process and extract text content from various document types."""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_doc_text,
            '.xlsx': self._extract_xlsx_text,
            '.xls': self._extract_xls_text,
            '.txt': self._extract_txt_text
        }
    
    async def process_document(self, document: DocumentMetadata, content: bytes) -> DocumentMetadata:
        """Process a document and extract text content."""
        
        file_type = document.file_type.lower()
        
        if file_type not in self.supported_types:
            logger.warning("Unsupported file type", file_type=file_type, file_name=document.file_name)
            return document
        
        try:
            # Extract text based on file type
            extractor = self.supported_types[file_type]
            extracted_text = await extractor(content)
            
            if extracted_text:
                # Update document with extracted content
                document.extracted_text = extracted_text
                document.content_preview = self._create_preview(extracted_text)
                
                # Extract additional metadata from content
                await self._extract_metadata_from_content(document, extracted_text)
                
                logger.info("Successfully processed document", 
                           file_name=document.file_name, 
                           text_length=len(extracted_text))
            else:
                logger.warning("No text extracted from document", file_name=document.file_name)
            
        except Exception as e:
            logger.error("Failed to process document", 
                        file_name=document.file_name, 
                        error=str(e))
        
        return document
    
    async def _extract_pdf_text(self, content: bytes) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error("Failed to extract PDF text", error=str(e))
            return None
    
    async def _extract_docx_text(self, content: bytes) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error("Failed to extract DOCX text", error=str(e))
            return None
    
    async def _extract_doc_text(self, content: bytes) -> Optional[str]:
        """Extract text from DOC file (legacy format)."""
        # Note: This is a simplified implementation
        # For production, consider using python-docx2txt or similar library
        logger.warning("DOC file processing is limited - consider converting to DOCX")
        return None
    
    async def _extract_xlsx_text(self, content: bytes) -> Optional[str]:
        """Extract text from XLSX file."""
        try:
            excel_file = io.BytesIO(content)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                
                text_parts.append("")  # Add blank line between sheets
            
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error("Failed to extract XLSX text", error=str(e))
            return None
    
    async def _extract_xls_text(self, content: bytes) -> Optional[str]:
        """Extract text from XLS file (legacy format)."""
        # Note: This would require xlrd library for older Excel formats
        logger.warning("XLS file processing not implemented - consider converting to XLSX")
        return None
    
    async def _extract_txt_text(self, content: bytes) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            logger.warning("Could not decode text file with common encodings")
            return None
            
        except Exception as e:
            logger.error("Failed to extract TXT text", error=str(e))
            return None
    
    def _create_preview(self, text: str, max_length: int = 500) -> str:
        """Create a preview of the extracted text."""
        if not text:
            return ""
        
        # Clean up the text
        cleaned_text = " ".join(text.split())
        
        if len(cleaned_text) <= max_length:
            return cleaned_text
        
        # Find a good breaking point near the max length
        preview = cleaned_text[:max_length]
        last_space = preview.rfind(' ')
        
        if last_space > max_length * 0.8:  # If we can find a space in the last 20%
            preview = preview[:last_space]
        
        return preview + "..."
    
    async def _extract_metadata_from_content(self, document: DocumentMetadata, text: str):
        """Extract metadata from document content using simple patterns."""
        
        text_lower = text.lower()
        
        # Extract potential client names (look for common patterns)
        client_patterns = [
            r'client:\s*([^\n\r]+)',
            r'for:\s*([^\n\r]+)',
            r'project owner:\s*([^\n\r]+)'
        ]
        
        import re
        for pattern in client_patterns:
            match = re.search(pattern, text_lower)
            if match:
                potential_client = match.group(1).strip()
                if len(potential_client) < 100:  # Reasonable length for client name
                    document.client_name = potential_client.title()
                    break
        
        # Extract potential project titles
        title_patterns = [
            r'project:\s*([^\n\r]+)',
            r'subject:\s*([^\n\r]+)',
            r'title:\s*([^\n\r]+)'
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, text_lower)
            if match:
                potential_title = match.group(1).strip()
                if len(potential_title) < 200:  # Reasonable length for title
                    document.project_title = potential_title.title()
                    break
        
        # Extract keywords (common engineering terms)
        engineering_keywords = [
            'seismic', 'structural', 'bridge', 'building', 'foundation',
            'concrete', 'steel', 'retrofit', 'design', 'analysis',
            'inspection', 'report', 'specification', 'drawing',
            'calculation', 'load', 'capacity', 'safety'
        ]
        
        found_keywords = []
        for keyword in engineering_keywords:
            if keyword in text_lower:
                found_keywords.append(keyword)
        
        document.keywords = found_keywords[:10]  # Limit to 10 keywords
        
        # Determine project status based on content
        if any(word in text_lower for word in ['final', 'completed', 'issued', 'approved']):
            document.project_status = 'completed'
        elif any(word in text_lower for word in ['draft', 'preliminary', 'review']):
            document.project_status = 'active'
        elif any(word in text_lower for word in ['on hold', 'suspended', 'paused']):
            document.project_status = 'on_hold'
