import asyncio
from typing import Optional, Dict, Any
import docx
import PyPDF2
import openpyxl
import io
import structlog
import extract_msg
import json
import csv
from PIL import Image

from src.models import DocumentMetadata

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """Process and extract text content from various document types."""
    
    def __init__(self):
        self.supported_types = {
            # Document formats
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_doc_text,
            '.txt': self._extract_txt_text,
            '.msg': self._extract_msg_text,
            
            # Spreadsheet formats
            '.xlsx': self._extract_xlsx_text,
            '.xls': self._extract_xls_text,
            '.csv': self._extract_csv_text,
            
            # Presentation formats
            '.pptx': self._extract_pptx_text,
            '.ppt': self._extract_ppt_text,
            
            # Image formats (with OCR potential)
            '.jpg': self._extract_image_metadata,
            '.jpeg': self._extract_image_metadata,
            '.png': self._extract_image_metadata,
            '.gif': self._extract_image_metadata,
            '.bmp': self._extract_image_metadata,
            '.tiff': self._extract_image_metadata,
            '.tif': self._extract_image_metadata,
            
            # CAD and drawing formats
            '.dwg': self._extract_cad_metadata,
            '.dxf': self._extract_dxf_text,
            '.dwf': self._extract_cad_metadata,
            '.svg': self._extract_svg_text,
            
            # Data formats
            '.json': self._extract_json_text,
            '.xml': self._extract_xml_text,
            
            # Archive formats
            '.zip': self._extract_archive_metadata,
            '.rar': self._extract_archive_metadata,
            '.7z': self._extract_archive_metadata
        }
    
    async def process_document(self, document: DocumentMetadata, content: bytes) -> DocumentMetadata:
        """Process a document and extract text content."""
        
        file_type = document.file_type.lower()
        
        if file_type not in self.supported_types:
            logger.warning("Unsupported file type", file_type=file_type, file_name=document.file_name)
            return document
        
        try:
            # Extract text based on file type
            extractor = self.supported_types[file_type]
            extracted_text = await extractor(content)
            
            if extracted_text:
                # Update document with extracted content
                document.extracted_text = extracted_text
                document.content_preview = self._create_preview(extracted_text)
                
                # Extract additional metadata from content
                await self._extract_metadata_from_content(document, extracted_text)
                
                logger.info("Successfully processed document", 
                           file_name=document.file_name, 
                           text_length=len(extracted_text))
            else:
                logger.warning("No text extracted from document", file_name=document.file_name)
            
        except Exception as e:
            logger.error("Failed to process document", 
                        file_name=document.file_name, 
                        error=str(e))
        
        return document
    
    async def _extract_pdf_text(self, content: bytes) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error("Failed to extract PDF text", error=str(e))
            return None
    
    async def _extract_docx_text(self, content: bytes) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error("Failed to extract DOCX text", error=str(e))
            return None
    
    async def _extract_doc_text(self, content: bytes) -> Optional[str]:
        """Extract text from DOC file (legacy format)."""
        # Note: This is a simplified implementation
        # For production, consider using python-docx2txt or similar library
        logger.warning("DOC file processing is limited - consider converting to DOCX")
        return None
    
    async def _extract_xlsx_text(self, content: bytes) -> Optional[str]:
        """Extract text from XLSX file."""
        try:
            excel_file = io.BytesIO(content)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                
                text_parts.append("")  # Add blank line between sheets
            
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error("Failed to extract XLSX text", error=str(e))
            return None
    
    async def _extract_xls_text(self, content: bytes) -> Optional[str]:
        """Extract text from XLS file (legacy format)."""
        # Note: This would require xlrd library for older Excel formats
        logger.warning("XLS file processing not implemented - consider converting to XLSX")
        return None
    
    async def _extract_txt_text(self, content: bytes) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            logger.warning("Could not decode text file with common encodings")
            return None
            
        except Exception as e:
            logger.error("Failed to extract TXT text", error=str(e))
            return None

    async def _extract_msg_text(self, content: bytes) -> Optional[str]:
        """Extract text from MSG (Outlook message) file."""
        try:
            import tempfile
            import os
            
            # MSG files need to be saved to disk for extract_msg to read them
            with tempfile.NamedTemporaryFile(suffix='.msg', delete=False) as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()
                
                try:
                    # Use extract_msg to parse the Outlook message
                    msg = extract_msg.Message(tmp_file.name)
                    
                    text_parts = []
                    
                    # Extract email metadata
                    if msg.subject:
                        text_parts.append(f"Subject: {msg.subject}")
                    
                    if msg.sender:
                        text_parts.append(f"From: {msg.sender}")
                    
                    if msg.to:
                        text_parts.append(f"To: {msg.to}")
                    
                    if msg.date:
                        text_parts.append(f"Date: {msg.date}")
                    
                    # Extract message body
                    if msg.body:
                        text_parts.append("\nMessage:")
                        text_parts.append(msg.body)
                    
                    # Extract attachments info (file names only, not content)
                    if hasattr(msg, 'attachments') and msg.attachments:
                        text_parts.append("\nAttachments:")
                        for attachment in msg.attachments:
                            if hasattr(attachment, 'longFilename') and attachment.longFilename:
                                text_parts.append(f"- {attachment.longFilename}")
                    
                    return "\n".join(text_parts) if text_parts else None
                    
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                        
        except Exception as e:
            logger.error("Failed to extract MSG text", error=str(e))
            return None

    async def _extract_csv_text(self, content: bytes) -> Optional[str]:
        """Extract text from CSV file."""
        try:
            import csv
            import io
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = content.decode(encoding)
                    csv_file = io.StringIO(text_content)
                    reader = csv.reader(csv_file)
                    
                    text_parts = []
                    row_count = 0
                    
                    for row in reader:
                        if row_count == 0:
                            text_parts.append("Headers: " + " | ".join(row))
                        else:
                            text_parts.append(" | ".join(row))
                        
                        row_count += 1
                        if row_count > 100:  # Limit to first 100 rows
                            text_parts.append(f"... and {row_count} more rows")
                            break
                    
                    return "\n".join(text_parts) if text_parts else None
                    
                except UnicodeDecodeError:
                    continue
                    
        except Exception as e:
            logger.error("Failed to extract CSV text", error=str(e))
            return None

    async def _extract_json_text(self, content: bytes) -> Optional[str]:
        """Extract text from JSON file."""
        try:
            import json
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1']
            
            for encoding in encodings:
                try:
                    text_content = content.decode(encoding)
                    data = json.loads(text_content)
                    
                    # Convert JSON to readable text
                    return json.dumps(data, indent=2, ensure_ascii=False)
                    
                except (UnicodeDecodeError, json.JSONDecodeError):
                    continue
                    
        except Exception as e:
            logger.error("Failed to extract JSON text", error=str(e))
            return None

    async def _extract_xml_text(self, content: bytes) -> Optional[str]:
        """Extract text from XML file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
                    
        except Exception as e:
            logger.error("Failed to extract XML text", error=str(e))
            return None

    async def _extract_image_metadata(self, content: bytes) -> Optional[str]:
        """Extract metadata and basic info from image files."""
        try:
            from PIL import Image
            import io
            
            image = Image.open(io.BytesIO(content))
            
            text_parts = []
            text_parts.append(f"Image Type: {image.format}")
            text_parts.append(f"Dimensions: {image.size[0]} x {image.size[1]} pixels")
            text_parts.append(f"Color Mode: {image.mode}")
            
            # Extract EXIF data if available
            if hasattr(image, '_getexif'):
                exif = image._getexif()
                if exif:
                    text_parts.append("EXIF Data:")
                    for key, value in exif.items():
                        if isinstance(value, str) and len(value) < 100:
                            text_parts.append(f"  {key}: {value}")
            
            # Note: For OCR text extraction from images, you'd need pytesseract
            # text_parts.append("Note: OCR text extraction requires pytesseract library")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error("Failed to extract image metadata", error=str(e))
            return None

    async def _extract_cad_metadata(self, content: bytes) -> Optional[str]:
        """Extract basic metadata from CAD files (.dwg, .dwf)."""
        try:
            # For now, extract basic file info
            # Full CAD parsing would require specialized libraries like ezdxf
            text_parts = []
            text_parts.append("CAD Drawing File")
            text_parts.append(f"File Size: {len(content)} bytes")
            
            # Check for common CAD text patterns in binary data
            content_str = str(content[:2000])  # First 2KB for patterns
            
            if 'AutoCAD' in content_str:
                text_parts.append("Software: AutoCAD")
            if 'DWG' in content_str:
                text_parts.append("Format: DWG")
                
            text_parts.append("Note: Full CAD text extraction requires specialized CAD libraries")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error("Failed to extract CAD metadata", error=str(e))
            return None

    async def _extract_dxf_text(self, content: bytes) -> Optional[str]:
        """Extract text from DXF files (ASCII CAD format)."""
        try:
            # DXF files are often text-based, so we can extract some content
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = content.decode(encoding)
                    
                    text_parts = []
                    text_parts.append("DXF CAD Drawing")
                    
                    # Extract layer names
                    lines = text_content.split('\n')
                    layers = set()
                    texts = []
                    
                    for i, line in enumerate(lines[:1000]):  # Limit scan
                        if line.strip() == '8':  # Layer code in DXF
                            if i + 1 < len(lines):
                                layer = lines[i + 1].strip()
                                if layer and not layer.isdigit():
                                    layers.add(layer)
                        
                        if line.strip() == '1':  # Text value code in DXF
                            if i + 1 < len(lines):
                                text = lines[i + 1].strip()
                                if text and len(text) > 1:
                                    texts.append(text)
                    
                    if layers:
                        text_parts.append("Layers: " + ", ".join(list(layers)[:10]))
                    
                    if texts:
                        text_parts.append("Text Content:")
                        text_parts.extend(texts[:20])  # First 20 text elements
                    
                    return "\n".join(text_parts)
                    
                except UnicodeDecodeError:
                    continue
                    
        except Exception as e:
            logger.error("Failed to extract DXF text", error=str(e))
            return None

    async def _extract_svg_text(self, content: bytes) -> Optional[str]:
        """Extract text from SVG files."""
        try:
            # SVG is XML-based, so decode as text
            encodings = ['utf-8', 'utf-16', 'latin-1']
            
            for encoding in encodings:
                try:
                    svg_content = content.decode(encoding)
                    
                    text_parts = []
                    text_parts.append("SVG Vector Drawing")
                    
                    # Extract text elements
                    import re
                    text_matches = re.findall(r'<text[^>]*>(.*?)</text>', svg_content, re.IGNORECASE | re.DOTALL)
                    
                    if text_matches:
                        text_parts.append("Text Elements:")
                        for text in text_matches[:20]:  # Limit to first 20
                            clean_text = re.sub(r'<[^>]+>', '', text).strip()
                            if clean_text:
                                text_parts.append(clean_text)
                    
                    return "\n".join(text_parts)
                    
                except UnicodeDecodeError:
                    continue
                    
        except Exception as e:
            logger.error("Failed to extract SVG text", error=str(e))
            return None

    async def _extract_pptx_text(self, content: bytes) -> Optional[str]:
        """Extract text from PowerPoint files."""
        try:
            # Note: This would require python-pptx library
            # For now, return placeholder
            return "PowerPoint Presentation - Text extraction requires python-pptx library"
            
        except Exception as e:
            logger.error("Failed to extract PPTX text", error=str(e))
            return None

    async def _extract_ppt_text(self, content: bytes) -> Optional[str]:
        """Extract text from legacy PowerPoint files."""
        try:
            return "Legacy PowerPoint Presentation - Text extraction requires specialized library"
            
        except Exception as e:
            logger.error("Failed to extract PPT text", error=str(e))
            return None

    async def _extract_archive_metadata(self, content: bytes) -> Optional[str]:
        """Extract file list from archive files."""
        try:
            import zipfile
            import io
            
            text_parts = []
            text_parts.append("Archive File Contents:")
            
            try:
                # Try as ZIP file
                with zipfile.ZipFile(io.BytesIO(content), 'r') as zip_file:
                    file_list = zip_file.namelist()
                    text_parts.append(f"Total files: {len(file_list)}")
                    text_parts.append("Files:")
                    
                    for filename in file_list[:50]:  # Limit to first 50 files
                        text_parts.append(f"  {filename}")
                    
                    if len(file_list) > 50:
                        text_parts.append(f"  ... and {len(file_list) - 50} more files")
                        
            except zipfile.BadZipFile:
                text_parts.append("Archive format not supported for content extraction")
                text_parts.append(f"File size: {len(content)} bytes")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error("Failed to extract archive metadata", error=str(e))
            return None
    
    def _create_preview(self, text: str, max_length: int = 500) -> str:
        """Create a preview of the extracted text."""
        if not text:
            return ""
        
        # Clean up the text
        cleaned_text = " ".join(text.split())
        
        if len(cleaned_text) <= max_length:
            return cleaned_text
        
        # Find a good breaking point near the max length
        preview = cleaned_text[:max_length]
        last_space = preview.rfind(' ')
        
        if last_space > max_length * 0.8:  # If we can find a space in the last 20%
            preview = preview[:last_space]
        
        return preview + "..."
    
    async def _extract_metadata_from_content(self, document: DocumentMetadata, text: str):
        """Extract metadata from document content using simple patterns."""
        
        text_lower = text.lower()
        
        # Extract potential client names (look for common patterns)
        client_patterns = [
            r'client:\s*([^\n\r]+)',
            r'for:\s*([^\n\r]+)',
            r'project owner:\s*([^\n\r]+)'
        ]
        
        import re
        for pattern in client_patterns:
            match = re.search(pattern, text_lower)
            if match:
                potential_client = match.group(1).strip()
                if len(potential_client) < 100:  # Reasonable length for client name
                    document.client_name = potential_client.title()
                    break
        
        # Extract potential project titles
        title_patterns = [
            r'project:\s*([^\n\r]+)',
            r'subject:\s*([^\n\r]+)',
            r'title:\s*([^\n\r]+)'
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, text_lower)
            if match:
                potential_title = match.group(1).strip()
                if len(potential_title) < 200:  # Reasonable length for title
                    document.project_title = potential_title.title()
                    break
        
        # Extract keywords (common engineering terms)
        engineering_keywords = [
            'seismic', 'structural', 'bridge', 'building', 'foundation',
            'concrete', 'steel', 'retrofit', 'design', 'analysis',
            'inspection', 'report', 'specification', 'drawing',
            'calculation', 'load', 'capacity', 'safety'
        ]
        
        found_keywords = []
        for keyword in engineering_keywords:
            if keyword in text_lower:
                found_keywords.append(keyword)
        
        document.keywords = found_keywords[:10]  # Limit to 10 keywords
        
        # Determine project status based on content
        if any(word in text_lower for word in ['final', 'completed', 'issued', 'approved']):
            document.project_status = 'completed'
        elif any(word in text_lower for word in ['draft', 'preliminary', 'review']):
            document.project_status = 'active'
        elif any(word in text_lower for word in ['on hold', 'suspended', 'paused']):
            document.project_status = 'on_hold'
