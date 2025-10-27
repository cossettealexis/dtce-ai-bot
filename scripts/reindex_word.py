#!/usr/bin/env python3
"""
Word Documents (.docx, .doc) reindexing script
Run this in parallel with other file type scripts
"""

import os
import sys
import asyncio
import io
from dotenv import load_dotenv
from azure.storage.blob import BlobServiceClient
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
import re
from datetime import datetime, timezone
import time
import docx
from openai import AsyncAzureOpenAI

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Load environment variables
load_dotenv()

def fast_extract_docx_content(blob_data: bytes, filename: str) -> str:
    """Fast DOCX extraction."""
    try:
        size_mb = len(blob_data) / (1024 * 1024)
        print(f"    📄 Extracting Word content ({size_mb:.1f}MB)...")
        
        docx_stream = io.BytesIO(blob_data)
        doc = docx.Document(docx_stream)
        
        text_content = ""
        paragraphs_processed = 0
        
        # Limit paragraphs for speed
        for paragraph in doc.paragraphs[:200]:
            if paragraph.text and paragraph.text.strip():
                text_content += paragraph.text.strip() + "\\n"
                paragraphs_processed += 1
                
                # Stop if we have enough content
                if len(text_content) > 100000:
                    break
        
        # Also extract from tables
        tables_processed = 0
        for table in doc.tables[:10]:  # Limit tables
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content += " | ".join(row_text) + "\\n"
                    tables_processed += 1
                    
                if len(text_content) > 12000:
                    break
            if len(text_content) > 12000:
                break
        
        result = text_content.strip()
        if result:
            print(f"    ✅ Word: {paragraphs_processed} paragraphs, {tables_processed} tables, {len(result)} chars")
            return result
        else:
            return "Word document (no text content found)"
            
    except Exception as e:
        return f"Word document (extraction failed: {str(e)[:100]})"

def fast_extract_doc_content(blob_data: bytes, filename: str) -> str:
    """Fast .doc (legacy) extraction - basic text recovery."""
    try:
        size_mb = len(blob_data) / (1024 * 1024)
        print(f"    📄 Extracting legacy Word content ({size_mb:.1f}MB)...")
        
        # Try to extract readable text from legacy .doc format
        text = blob_data.decode('utf-8', errors='ignore')
        
        # Clean up the text - remove control characters and binary junk
        import re
        text = re.sub(r'[\\x00-\\x08\\x0b\\x0c\\x0e-\\x1f\\x7f-\\xff]', ' ', text)
        text = re.sub(r'\\s+', ' ', text)
        
        # Extract words that look like real text
        words = text.split()
        filtered_words = []
        
        for word in words:
            # Keep words that are mostly alphabetic and reasonable length
            if (len(word) >= 3 and len(word) <= 50 and 
                sum(1 for c in word if c.isalpha()) / len(word) > 0.7):
                filtered_words.append(word)
                
                if len(filtered_words) >= 1000:  # Limit for speed
                    break
        
        if len(filtered_words) > 20:
            result = ' '.join(filtered_words)
            print(f"    ✅ Legacy Word: {len(filtered_words)} words, {len(result)} chars")
            return result
        else:
            return "Legacy Word document (limited text extraction)"
            
    except Exception as e:
        return f"Legacy Word document (extraction failed: {str(e)[:100]})"

def document_needs_update(blob, existing_docs: dict) -> bool:
    """Check if Word document needs updating."""
    document_id = re.sub(r'[^a-zA-Z0-9_-]', '_', blob.name)
    document_id = re.sub(r'_+', '_', document_id).strip('_')
    
    if document_id not in existing_docs:
        return True
    
    existing_doc = existing_docs[document_id]
    existing_content = existing_doc.get('content', '')
    
    # Update if content is bad
    if (not existing_content or 
        len(existing_content) < 50 or
        'no text content found' in existing_content.lower() or
        'extraction failed' in existing_content.lower()):
        return True
    
    # Update if blob is newer
    try:
        existing_modified = existing_doc.get('last_modified', '')
        if existing_modified and blob.last_modified:
            existing_time = datetime.fromisoformat(existing_modified.replace('Z', '+00:00'))
            if blob.last_modified.replace(tzinfo=timezone.utc) > existing_time:
                return True
    except:
        pass
    
    return False

async def generate_embeddings(openai_client: AsyncAzureOpenAI, text: str) -> list:
    """Generate embeddings with truncation."""
    try:
        truncated_text = text[:4000]
        response = await openai_client.embeddings.create(
            model="text-embedding-3-small",
            input=truncated_text
        )
        return response.data[0].embedding
    except:
        return []

async def process_word_docs():
    """Process only Word documents."""
    print("📝 WORD DOCUMENTS REINDEXING SCRIPT")
    print("=" * 50)
    
    # Initialize clients
    connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    search_endpoint = os.getenv("AZURE_SEARCH_SERVICE_ENDPOINT")
    search_key = os.getenv("AZURE_SEARCH_API_KEY")
    index_name = os.getenv("AZURE_SEARCH_INDEX_NAME", "dtce-documents-index")
    container_name = os.getenv("AZURE_STORAGE_CONTAINER_NAME", "dtce-documents")
    
    storage_client = BlobServiceClient.from_connection_string(connection_string)
    search_client = SearchClient(
        endpoint=search_endpoint,
        index_name=index_name,
        credential=AzureKeyCredential(search_key)
    )
    
    openai_client = AsyncAzureOpenAI(
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version="2024-02-15-preview",
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
    )
    
    container_client = storage_client.get_container_client(container_name)
    
    # Get existing documents
    print("📋 Loading existing Word docs from search index...")
    existing_docs = {}
    try:
        results = search_client.search("*", select=["id", "last_modified", "content"], top=10000)
        for doc in results:
            if doc.get('id'):
                existing_docs[doc['id']] = {
                    'last_modified': doc.get('last_modified'),
                    'content': doc.get('content', '')
                }
        print(f"✅ Loaded {len(existing_docs)} existing documents")
    except Exception as e:
        print(f"⚠️ Could not load existing docs: {e}")
    
    # Find Word files that need updating
    print("🔍 Scanning for Word documents...")
    word_blobs = []
    word_extensions = ['.docx', '.dotx', '.doc', '.dot']
    
    # Projects Word docs first
    try:
        for blob in container_client.list_blobs(name_starts_with="Projects/"):
            if any(blob.name.lower().endswith(ext) for ext in word_extensions):
                if document_needs_update(blob, existing_docs):
                    word_blobs.append(blob)
    except Exception as e:
        print(f"⚠️ Error scanning Projects Word docs: {e}")
    
    # Other Word docs
    try:
        for blob in container_client.list_blobs():
            if (any(blob.name.lower().endswith(ext) for ext in word_extensions) and 
                not blob.name.startswith("Projects/")):
                if document_needs_update(blob, existing_docs):
                    word_blobs.append(blob)
    except Exception as e:
        print(f"⚠️ Error scanning other Word docs: {e}")
    
    total_docs = len(word_blobs)
    print(f"📊 Found {total_docs} Word documents to process")
    
    if total_docs == 0:
        print("✅ All Word documents are up to date!")
        return
    
    # Process Word documents
    processed = 0
    errors = 0
    start_time = time.time()
    
    for i, blob in enumerate(word_blobs, 1):
        try:
            print(f"\\n[{i}/{total_docs}] Processing: {blob.name}")
            
            # Download blob
            blob_client = storage_client.get_blob_client(container=container_name, blob=blob.name)
            blob_data = blob_client.download_blob().readall()
            
            # Extract content based on file type
            filename_lower = blob.name.lower()
            if filename_lower.endswith(('.docx', '.dotx')):
                content = fast_extract_docx_content(blob_data, blob.name)
            else:  # .doc, .dot
                content = fast_extract_doc_content(blob_data, blob.name)
            
            # Extract metadata
            folder_path = blob.name.rsplit('/', 1)[0] if '/' in blob.name else ''
            filename = os.path.basename(blob.name)
            
            # Project info
            project_name = ""
            year = None
            if folder_path and '/Projects/' in f"/{folder_path}/":
                parts = folder_path.split('/')
                if 'Projects' in parts:
                    idx = parts.index('Projects')
                    if idx + 1 < len(parts):
                        project_name = parts[idx + 1]
                    if idx + 2 < len(parts):
                        try:
                            year = int(parts[idx + 2])
                        except:
                            pass
            
            # Generate embeddings
            content_vector = await generate_embeddings(openai_client, content)
            
            # Create document
            document_id = re.sub(r'[^a-zA-Z0-9_-]', '_', blob.name)
            document_id = re.sub(r'_+', '_', document_id).strip('_')
            
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if filename_lower.endswith(('.doc', '.dot')):
                content_type = "application/msword"
            
            search_document = {
                "id": document_id,
                "blob_name": blob.name,
                "blob_url": blob_client.url,
                "filename": filename,
                "content_type": content_type,
                "folder": folder_path,
                "size": blob.size or 0,
                "content": content,
                "content_vector": content_vector,
                "last_modified": blob.last_modified.isoformat(),
                "created_date": blob.creation_time.isoformat() if blob.creation_time else blob.last_modified.isoformat(),
                "project_name": project_name,
                "year": year
            }
            
            # Upload to search
            result = search_client.upload_documents([search_document])
            if result[0].succeeded:
                processed += 1
                print(f"    ✅ Indexed successfully")
            else:
                errors += 1
                print(f"    ❌ Index failed: {result[0].error_message}")
            
            # Progress
            if i % 5 == 0:
                elapsed = time.time() - start_time
                rate = processed / elapsed * 60
                remaining = total_docs - i
                eta = remaining / rate if rate > 0 else 0
                print(f"    📊 Progress: {processed}/{total_docs}, {rate:.1f} docs/min, ETA: {eta:.1f}min")
                
        except Exception as e:
            errors += 1
            print(f"    ❌ Error: {str(e)[:100]}")
    
    # Summary
    elapsed = time.time() - start_time
    rate = processed / elapsed * 60 if elapsed > 0 else 0
    
    print(f"\\n🎉 WORD PROCESSING COMPLETE!")
    print(f"✅ Processed: {processed}")
    print(f"❌ Errors: {errors}")
    print(f"⏱️ Time: {elapsed/60:.1f} minutes")
    print(f"📈 Rate: {rate:.1f} docs/minute")

if __name__ == "__main__":
    asyncio.run(process_word_docs())
